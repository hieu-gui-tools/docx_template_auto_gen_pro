"""Parser để tìm và xử lý các trường {{ten_truong|type}} trong DOCX template."""
from __future__ import annotations

import importlib.util
import inspect
import re
import sys
from pathlib import Path
from typing import Any

import docx

# ── Regex ────────────────────────────────────────────────────────────────────
FIELD_PATTERN = re.compile(r"\{\{([^|{}]+)(?:\|([^{}]+))?\}\}")

# Các type được hỗ trợ
VALID_TYPES = {
    "text", "m_text", "text_lower", "text_upper", "text_capitalize", "text_title",
    "number", "integer", "float",
    "date", "datetime",
    "script",
}


def _runs_full_text(para) -> str:
    return "".join(r.text for r in para.runs)


def extract_fields(docx_path: str | Path) -> list[dict]:
    """
    Trả về list các dict: {name, type, para_idx, run_indices, raw_tag}
    Mỗi phần tử là một trường phát hiện được trong template.
    """
    doc = docx.Document(str(docx_path))
    fields: list[dict] = []
    seen: set[str] = set()

    def _scan_text(text: str) -> list[dict]:
        res = []
        for m in FIELD_PATTERN.finditer(text):
            name = m.group(1).strip()
            ftype_raw = m.group(2).strip() if m.group(2) else ""
            if not ftype_raw or ftype_raw.lower() not in VALID_TYPES:
                ftype = "text"
                is_default = True
            else:
                ftype = ftype_raw
                is_default = False
            res.append({
                "name": name,
                "type": ftype,
                "raw": m.group(0),
                "is_default": is_default,
                "original_type": ftype_raw
            })
        return res

    # Scan paragraphs
    for para in doc.paragraphs:
        full = _runs_full_text(para)
        for field_dict in _scan_text(full):
            name = field_dict["name"]
            ftype = field_dict["type"]
            key = f"{name}|{ftype}"
            if key not in seen:
                seen.add(key)
                fields.append(field_dict)

    # Scan tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full = _runs_full_text(para)
                    for field_dict in _scan_text(full):
                        name = field_dict["name"]
                        ftype = field_dict["type"]
                        key = f"{name}|{ftype}"
                        if key not in seen:
                            seen.add(key)
                            fields.append(field_dict)

    # Scan headers/footers
    for section in doc.sections:
        for hdr in (section.header, section.footer):
            if hdr is not None:
                for para in hdr.paragraphs:
                    full = _runs_full_text(para)
                    for field_dict in _scan_text(full):
                        name = field_dict["name"]
                        ftype = field_dict["type"]
                        key = f"{name}|{ftype}"
                        if key not in seen:
                            seen.add(key)
                            fields.append(field_dict)

    return fields


def _replace_in_para(para, replacements: dict[str, str]) -> None:
    """Thay thế placeholder trong một paragraph, giữ nguyên format run đầu tiên."""
    full = _runs_full_text(para)
    new_full = full
    for tag, value in replacements.items():
        new_full = new_full.replace(tag, value)
    if new_full == full:
        return
    # Ghi vào run đầu, xóa run còn lại
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_full)


def generate_document(
    docx_path: str | Path,
    field_values: dict[str, str],    # {raw_tag: formatted_value}
    output_path: str | Path,
) -> Path:
    """Tạo file mới từ template với các trường đã được thay thế."""
    doc = docx.Document(str(docx_path))

    def _process_para(para):
        _replace_in_para(para, field_values)

    for para in doc.paragraphs:
        _process_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_para(para)

    for section in doc.sections:
        for hdr in (section.header, section.footer):
            if hdr is not None:
                for para in hdr.paragraphs:
                    _process_para(para)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


# ── Script type ──────────────────────────────────────────────────────────────

def load_script_module(docx_path: str | Path):
    """Load file .py cùng tên template."""
    py_path = Path(docx_path).with_suffix(".py")
    if not py_path.exists():
        return None
    spec = importlib.util.spec_from_file_location("_tpl_script", py_path)
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)
    return mod


def get_script_args(mod, func_name: str) -> list[str]:
    """Trả về list tên các argument của hàm (ngoài *args, **kwargs không tên)."""
    fn = getattr(mod, func_name, None)
    if fn is None:
        return []
    sig = inspect.signature(fn)
    return [
        p.name for p in sig.parameters.values()
        if p.kind in (
            inspect.Parameter.POSITIONAL_OR_KEYWORD,
            inspect.Parameter.KEYWORD_ONLY,
        )
    ]

def get_script_args_info(mod, func_name: str) -> dict[str, Any]:
    """Trả về dict name -> annotation của argument."""
    fn = getattr(mod, func_name, None)
    if fn is None:
        return {}
    sig = inspect.signature(fn)
    res = {}
    for p in sig.parameters.values():
        if p.kind in (inspect.Parameter.POSITIONAL_OR_KEYWORD, inspect.Parameter.KEYWORD_ONLY):
            res[p.name] = p.annotation
    return res


def call_script_func(mod, func_name: str, kwargs: dict[str, Any]) -> Any:
    """Gọi hàm script và trả về kết quả."""
    fn = getattr(mod, func_name, None)
    if fn is None:
        raise ValueError(f"Không tìm thấy hàm '{func_name}' trong script.")
    result = fn(**kwargs) if kwargs else fn()
    return result
